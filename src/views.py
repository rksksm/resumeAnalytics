# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from models import JD
from django.shortcuts import render
from django.conf import settings
from django.core.files.storage import FileSystemStorage
import RAKE
import nltk, string
from sklearn.feature_extraction.text import TfidfVectorizer
from docx import Document
import unicodedata
import csv
import pandas as pd
import itertools
import numpy as np
import os
from docx import Document
import unicodedata
from tabulate import tabulate
from fuzzywuzzy import fuzz
from fuzzywuzzy import process
import MySQLdb
import mysql.connector

# Create your views here.


stemmer = nltk.stem.porter.PorterStemmer()
remove_punctuation_map = dict((ord(char), None) for char in string.punctuation)
rake_object = RAKE.Rake(RAKE.SmartStopList())


def convertFormat(text):
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore')


def stem_tokens(tokens):
    return [stemmer.stem(item) for item in tokens]


def normalize(text):
    return stem_tokens(nltk.word_tokenize(text.lower().translate(remove_punctuation_map)))


def cosine_sim(text1, text2):
    try:
        vectorizer = TfidfVectorizer(tokenizer=normalize, stop_words='english')
        tfidf = vectorizer.fit_transform([text1, text2])
        return ((tfidf * tfidf.T).A)[0, 1]
    except Exception as e:
        return 0


def getKeywords(text):
    return rake_object.run(text)


def convertFormat(text):
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore')


def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def removeStopWords(filename):

    fullText = []
    for para in filename.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def fetchData(database, host, port, user, password, sql):
    con = mysql.connector.connect(host=host, port=port, database=database, user=user, password=password)
    cursor = con.cursor()
    cursor.execute(sql)
    data = cursor.fetchall()
    return data


def insertData(database, host, port, user, password, sql):
    db = MySQLdb.connect(host=host, port=port, db=database, user=user, passwd=password)
    cursor = db.cursor()
    cursor.execute(sql)
    db.commit()
    return "Data Inserted"


# filename = "/home/netlink/Downloads/ds_resume.docx"
def index(request):
    if request.method == 'POST' and request.FILES['myfile']:
        myfile = request.FILES['myfile']
        fs = FileSystemStorage()
        filename = fs.save(myfile.name, myfile)
        uploaded_file_url = fs.url(filename)
        print os.getcwd()+"/"+uploaded_file_url

        jd = JD.objects.filter(isActive='True')
        context = {'JD_title': [q.title for q in jd], 'JD_description': [q.description for q in jd]}
        selectedJD = []
        for i in range(0,len(context['JD_title'])):
            tempDict = {}
            text = convertFormat(getText(os.getcwd()+"/"+uploaded_file_url.replace("%20"," ")))
            text = getKeywords(text)
            resume = ' '.join(zip(*text)[0])
            text1 = getKeywords(context['JD_description'][i])
            jd = ' '.join(zip(*text1)[0])
            a = fuzz.token_set_ratio(resume, jd)
            b = fuzz.ratio(resume, jd)
            c = fuzz.token_sort_ratio(resume, jd)
            d = fuzz.partial_ratio(resume, jd)
            if a>50:
                tempDict["title"] = context['JD_title'][i]
                tempDict["description"] = context['JD_description'][i]
                selectedJD.append(tempDict)
                print context['JD_title'][i]
        if selectedJD == []:
            tempDict = {}
            tempDict["title"] = "No Title Found"
            tempDict["description"] = "No Description Found"
            selectedJD.append(tempDict)
        return render(request, 'simple_upload.html', {
            'uploaded_file_url': uploaded_file_url, 'selectedJD':selectedJD
        })
    return render(request, 'simple_upload.html')
